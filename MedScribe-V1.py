import sys
import sounddevice as sd
import numpy as np
import scipy.io.wavfile as wav
import time
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QLabel, QTextEdit, QVBoxLayout, QHBoxLayout, QWidget, QMessageBox, QLineEdit, QComboBox, QCheckBox, QInputDialog
from PyQt5.QtCore import QTimer, Qt
from PyQt5.QtGui import QIcon, QFont, QPixmap
from openai import OpenAI
from datetime import datetime
import os
import shutil
from docx import Document
import math
from scipy.io import wavfile
import uuid

def get_mac_address():
    mac = hex(uuid.getnode()).replace('0x', '')
    return ':'.join(mac[i:i + 2] for i in range(0, len(mac), 2))

client = OpenAI(api_key='<OpenAI API Key>')

# Global variables
recording = False
paused = False
audio_data = []  # List to store audio data during multiple recording sessions
sample_rate = 44100
filename = None
stream = None
start_time = None
stop = False
orginal_text = ""
max_audio_files = 10
#base_dir = os.path.dirname(sys.executable)
#base_dir = '/Users/kavyasreesoma/Documents/VoiceProcessing'

if getattr(sys, 'frozen', False):
    # If the app is running as a bundled executable (e.g., in a .app on macOS)
    base_dir = os.path.dirname(sys.executable)
    
    # For macOS apps, navigate out of the 'MacOS' folder to the base directory of the app bundle
    if sys.platform == 'darwin':
        base_dir = os.path.abspath(os.path.join(base_dir, '..', '..', '..'))
else:
    # If running as a regular Python script
    base_dir = os.path.dirname(os.path.abspath('/Users/kavyasreesoma/Documents/VoiceProcessing/Untitled1.ipynb'))

audio_folder_path = base_dir + '/audioFiles'
sn_path = base_dir + '/SOAP_Notes'

if not os.path.exists(audio_folder_path):
    os.makedirs(audio_folder_path)
if not os.path.exists(sn_path):
    os.makedirs(sn_path)

def split_wav_file(file_path, chunk_length_ms=90*1000):
    global base_dir
    output_dir = base_dir + '/temp'
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    else:
        shutil.rmtree(output_dir)
        os.makedirs(output_dir)
    
    # Read the wav file using scipy.io.wavfile
    sample_rate, audio_data = wavfile.read(file_path)
    
    # Calculate total length of the audio in milliseconds
    total_length_ms = (len(audio_data) / sample_rate) * 1000
    
    # Calculate the number of chunks
    chunk_length_samples = int((chunk_length_ms / 1000) * sample_rate)
    number_of_chunks = math.ceil(len(audio_data) / chunk_length_samples)
    
    l=[]
    # Split the audio into chunks and save each one
    for i in range(number_of_chunks):
        start_sample = i * chunk_length_samples
        end_sample = min((i + 1) * chunk_length_samples, len(audio_data))
        chunk_data = audio_data[start_sample:end_sample]
        
        # Define the output file name
        output_file_path = os.path.join(output_dir, f"chunk_{i + 1}.wav")
        l.append(output_file_path)
        
        # Export the chunk as a wav file
        wavfile.write(output_file_path, sample_rate, chunk_data)
        
    return l
        
# Function to get the length of the recorded audio in seconds
def get_audio_length():
    global audio_data, sample_rate
    if audio_data:
        total_frames = sum(chunk.shape[0] for chunk in audio_data)
        return total_frames / sample_rate
    return 0

# Function to save the recorded audio as a .wav file
def save_recording(filename_wav):
    global audio_data, sample_rate
    if audio_data:
        audio_array = np.concatenate(audio_data, axis=0)
        wav.write(filename_wav, sample_rate, audio_array)
        audio_data.clear()
        

# Function to limit audio files to 10
def limit_audio_files(directory, max_files):
    files = sorted([os.path.join(directory, f) for f in os.listdir(directory) if f.endswith(".wav")], key=os.path.getmtime)
    if len(files) >= max_files:
        # Delete the oldest file
        os.remove(files[0])

# Function to append audio chunks to the list
def audio_callback(indata, frames, time, status):
    global audio_data
    audio_data.append(indata.copy())

# Main application window
class VoiceRecorderApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Voice Recorder")
        self.showFullScreen()
#         self.setGeometry(100, 100, 900, 600)
        self.system_prompt = "You are AI medical scriber. you will be provided with a conversation between doctor and patient by user. Understand conversation by going through it twice and give output in SOAP format"


        # Custom font and stylesheet
        self.setStyleSheet("""
            QMainWindow {
                background-color: #e6e6e6;
            }
            QPushButton {
                font-size: 16px;
                font-family: Arial;
                padding: 10px;
                color: white;
                background-color: #4CAF50;
                border: none;
                border-radius: 10px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
             QPushButton:disabled {
                background-color: #777878;
            }
            
            QLabel {
                font-size: 18px;
                font-weight: bold;
                color: #333333;
            }
            QTextEdit {
                background-color: #f9f9f9;
                font-family: Arial;
                font-size: 14px;
                border: 2px solid #dddddd;
                border-radius: 10px;
            }
            QLineEdit {
                background-color: #f9f9f9;
                font-family: Arial;
                font-size: 14px;
                border: 2px solid #dddddd;
                border-radius: 10px;
            }
            QComboBox {
                font-size: 16px;
                font-family: Arial;
                padding: 5px;
            }
        """)

        # Load icons for buttons
        self.start_icon = QIcon(os.path.join(sys._MEIPASS, "icons/start.png"))
        self.pause_icon = QIcon(os.path.join(sys._MEIPASS, "icons/pause.png"))
        self.resume_icon = QIcon(os.path.join(sys._MEIPASS, "icons/resume.png"))
        self.stop_icon = QIcon(os.path.join(sys._MEIPASS, "icons/stop.png"))

        # Timer to update recording time
        self.timer = QTimer()
        self.timer.timeout.connect(self.update_timer)

        # Recording timer label
        self.timer_label = QLabel("Recording Time: 0s")
        self.timer_label.setAlignment(Qt.AlignLeft)
        self.timer_label.setFont(QFont("Arial", 16))

        # Buttons with icons
        self.start_button = QPushButton("Start Recording", self)
        self.start_button.setIcon(self.start_icon)
        self.start_button.clicked.connect(self.start_recording)
        self.start_button.setFixedWidth(150)

        self.pause_resume_button = QPushButton("Pause Recording", self)
        self.pause_resume_button.setIcon(self.pause_icon)
        self.pause_resume_button.setEnabled(False)
        self.pause_resume_button.clicked.connect(self.pause_resume_recording)
        self.pause_resume_button.setFixedWidth(170)

        self.stop_button = QPushButton("Stop Recording", self)
        self.stop_button.setIcon(self.stop_icon)
        self.stop_button.setEnabled(False)
        self.stop_button.clicked.connect(self.stop_recording)
        self.stop_button.setFixedWidth(150)
        
        
        # Doctor and Patient labels and text boxes
        self.doctor_label = QLabel("Doctor:")
        self.doctor_label.setFont(QFont("Arial", 14))
        
        self.doctor_name_label = QLabel("<Doctor Name>")  # Default doctor name
        self.doctor_name_label.setFont(QFont("Arial", 14))
        
        self.patient_label = QLabel("Patient:")
        self.patient_label.setFont(QFont("Arial", 14))
        
#         self.patient_text = QLineEdit()  # Patient name as editable text field
#         self.patient_text.setFont(QFont("Arial", 12))
#         self.patient_text.setFixedHeight(30)
#         self.patient_text.setFixedWidth(200)


        # Initially patient name will be blank and only the text box is visible
        self.patient_name_label = QLabel()  # Label for patient name, hidden initially
        self.patient_name_label.setFont(QFont("Arial", 12))
#         self.patient_name_label.hide()  # Hide the label initially
        self.patient_name_label.setFixedHeight(30)
        self.patient_name_label.setFixedWidth(200)

#         self.edit_button = QPushButton("Confirm")  # Button to toggle between edit and confirm
#         self.edit_button.setFont(QFont("Arial", 12))
#         self.edit_button.clicked.connect(self.toggle_edit_confirm)
        
        # Text areas for notes
        self.text_area1 = QTextEdit()
        self.text_area1.setReadOnly(True)
        self.text_area1.setText("Start the recording or select the existing audio!")
        self.text_area1.setFont(QFont("Arial", 12))

        self.text_area2 = QTextEdit()
        self.text_area2.setReadOnly(True)
        self.text_area2.setText("Your SOAP note will be generated here :)")
        self.text_area2.setFont(QFont("Arial", 12))
        
        self.edit_text_area_2 = QPushButton("Edit")  # Button to toggle between edit and confirm
        self.edit_text_area_2.setFont(QFont("Arial", 12))
        self.edit_text_area_2.clicked.connect(self.toggle_edit_save)
        
        self.download = QPushButton("Download")  # Button to toggle between edit and confirm
        self.download.setFont(QFont("Arial", 12))
        self.download.clicked.connect(self.on_click_download)

        self.checkbox = QCheckBox("Enable S.O.A.P. generation for older recordings")  # You can set any text for the checkbox
        self.checkbox.setFont(QFont("Arial", 14))
        self.checkbox.stateChanged.connect(self.toggle_picklist_visibility)
        
        self.picklist_label = QLabel("Select an Audio File:")
        self.picklist_label.setFont(QFont("Arial", 14))
        
        self.wav_file_picklist = QComboBox()
        self.wav_file_picklist.setFont(QFont("Arial", 12))
        self.wav_file_picklist.setFixedWidth(250)
        self.populate_picklist()
        
        self.generate_button = QPushButton("Generate S.O.A.P Notes")  # Button to toggle between edit and confirm
        self.generate_button.setFont(QFont("Arial", 12))
        self.generate_button.clicked.connect(self.OnClick_Generate)
        
        checkbox_layout = QHBoxLayout()
        checkbox_layout.addWidget(self.checkbox)
        checkbox_layout.setContentsMargins(0, 0, 0, 0)
        
        picklist_layout = QHBoxLayout()
        picklist_layout.addWidget(self.picklist_label)
        picklist_layout.addWidget(self.wav_file_picklist)
        picklist_layout.addWidget(self.generate_button)
        picklist_layout.addStretch(1)
        #picklist_layout.setContentsMargins(0, 0, 120, 0)
        
        # Wrap the picklist_layout in a QWidget container
        self.picklist_container = QWidget()
        self.picklist_container.setLayout(picklist_layout)
        self.picklist_container.setVisible(False)
        
        doc_layout = QHBoxLayout()
        doc_layout.addWidget(self.doctor_label)
        doc_layout.addWidget(self.doctor_name_label)
        doc_layout.addStretch(1)
        
        
        # Layout for patient name and edit/confirm button
        patient_layout = QHBoxLayout()  # Add patient layout back
        patient_layout.addWidget(self.patient_label)
#         patient_layout.addWidget(self.patient_text)  # Text field is visible initially
        patient_layout.addWidget(self.patient_name_label)  # Label is hidden initially
#         patient_layout.addWidget(self.edit_button)
        patient_layout.addStretch(1)

        
        # Layout for buttons
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.start_button)
        button_layout.addWidget(self.pause_resume_button)
        button_layout.addWidget(self.stop_button)
        
        
        recorder_layout = QVBoxLayout()
        recorder_layout.addWidget(self.timer_label)
        recorder_layout.addLayout(button_layout)
        
        info_layout = QVBoxLayout()
        info_layout.addLayout(doc_layout)
        info_layout.addLayout(patient_layout)
        info_layout.setContentsMargins(240, 0, 0, 0)

        # Combine buttons and doctor/patient info side by side

        controls_layout = QHBoxLayout()
        controls_layout.addLayout(recorder_layout)
        controls_layout.addLayout(info_layout)
        #controls_layout.addLayout(picklist_layout)
        
        options_layout = QVBoxLayout()
        options_layout.addLayout(controls_layout)
        options_layout.addLayout(checkbox_layout)
        options_layout.addWidget(self.picklist_container)

        edit_download_layout = QHBoxLayout()
        edit_download_layout.addWidget(self.edit_text_area_2)
        edit_download_layout.addWidget(self.download)
        
        text2_layout = QVBoxLayout()
        text2_layout.addWidget(self.text_area2)
        text2_layout.addLayout(edit_download_layout)

        # Layout for text areas
        text_layout = QHBoxLayout()
        text_layout.addWidget(self.text_area1)
        text_layout.addLayout(text2_layout)
        

        # Main layout
        main_layout = QVBoxLayout()
        #main_layout.addWidget(self.timer_label)
        main_layout.addLayout(options_layout)
        main_layout.addLayout(text_layout)

        container = QWidget()
        container.setLayout(main_layout)
        self.setCentralWidget(container)
        
    def populate_picklist(self):
        """Populate the picklist with .wav files from the audio folder."""
        self.wav_file_picklist.clear()
        self.wav_file_picklist.insertItem(0, "Select audio file")
        self.wav_file_picklist.setCurrentIndex(0)
        self.wav_file_picklist.model().item(0).setEnabled(False)
        if os.path.exists(audio_folder_path):
            # Get all .wav files in the folder
            wav_files = [f for f in os.listdir(audio_folder_path) if f.endswith('.wav')]
            # Add each .wav file to the combobox
            for wav_file in wav_files:
                self.wav_file_picklist.addItem(wav_file)
        else:
            self.wav_file_picklist.addItem("No files found")  # Fallback if folder doesn't exist

    def toggle_edit_save(self):
        if self.edit_text_area_2.text() == 'Edit':
            self.download.setEnabled(False)
            self.text_area2.setReadOnly(False)
            self.text_area2.setStyleSheet("background-color: #f3f3f3;")
            self.edit_text_area_2.setText("Save")
        else:
            self.download.setEnabled(True)
            self.text_area2.setReadOnly(True)
            self.text_area2.setStyleSheet("background-color: #f9f9f9;")
            self.edit_text_area_2.setText("Edit")
            
    def on_click_download(self):
        if self.text_area2.toPlainText() == "Your SOAP note will be generated here :)":
            QMessageBox.warning(self, "Warning", "No SOAP note is available to download!")
            return
        if self.checkbox.isChecked():
            audio_file = self.wav_file_picklist.currentText()
            if audio_file == "Select audio file":
                QMessageBox.warning(self, "Warning", "Please select the audio file")
                return
            patient_name = audio_file.split('-')[0]
#         else:
#             if not self.patient_text.text().strip():
#                 QMessageBox.warning(self, "Warning", "Please enter the patient's name")
#                 return
#             patient_name = self.patient_text.text().replace(" ", "_")
        
        doc = Document()
        doc.add_heading('S.O.A.P. Note', 0)
        doc.add_paragraph(self.text_area2.toPlainText())
        
        
        date_str = datetime.now().strftime("%Y-%m-%d")
        filename_docx = f"{patient_name}-{date_str}.docx"
        file_path = os.path.join(sn_path, filename_docx)
        
        doc.save(file_path)
#         self.patient_name_label.setText('')
        
        QMessageBox.information(self, "Info", f"S.O.A.P. Note saved as {filename_docx}")

    def show_input_dialog(self):
        # Show the input dialog and collect the input
        text, ok = QInputDialog.getText(self, "Patient Details", "Enter the Patient name:")

        if ok and text:  # Check if the user pressed OK and entered something
            self.patient_name_label.setText(text)

    # Start recording
    def start_recording(self):
        global recording, paused, audio_data, stream, start_time, stop
        self.show_input_dialog()
        if self.checkbox.isChecked():
            self.checkbox.setCheckState(Qt.Unchecked)
            #self.toggle_picklist_visibility(Qt.Unchecked)
        if not self.patient_name_label.text().strip():
            QMessageBox.warning(self, "Warning", "Please enter the patient's name before starting the recording.")
            return
        self.text_area2.setText("Your SOAP note will be generated here :)")
        self.text_area1.setText("Audio recording in progress...")
#         if self.edit_button.text() == "Confirm":
#             self.toggle_edit_confirm()
        if not recording:
            recording = True
            paused = False
            stop = False
            audio_data.clear()
            start_time = time.time()
            self.timer.start(1000)
            self.start_button.setEnabled(False)
            self.pause_resume_button.setEnabled(True)
            self.stop_button.setEnabled(True)
            self.checkbox.setEnabled(False)

            stream = sd.InputStream(samplerate=sample_rate, channels=1, callback=audio_callback)
            stream.start()

    # Pause or resume recording
    def pause_resume_recording(self):
        global paused, stream
        if recording and not paused:
            self.text_area1.setText("Audio recording PAUSED")
            paused = True
            stream.stop()
            self.pause_resume_button.setText("Resume Recording")
            self.pause_resume_button.setIcon(self.resume_icon)
        elif recording and paused:
            self.text_area1.setText("Audio recording in progress...")
            paused = False
            stream.start()
            self.pause_resume_button.setText("Pause Recording")
            self.pause_resume_button.setIcon(self.pause_icon)

    # Stop recording
    def stop_recording(self):
        global recording, audio_data, stream, stop, filename_wav, system_prompt
        self.pause_resume_button.setText("Pause Recording")
        self.pause_resume_button.setIcon(self.pause_icon)
        if recording:
            recording = False
            paused = False
            stop = True
            stream.stop()
            stream.close()
            self.timer.stop()
            self.start_button.setEnabled(True)
            self.pause_resume_button.setEnabled(False)
            self.stop_button.setEnabled(False)
            self.checkbox.setEnabled(True)
            
            patient_name = self.patient_name_label.text().replace(" ", "_")
            date_str = datetime.now().strftime("%Y-%m-%d")
            filename_wav = f"{patient_name}-{date_str}.wav"
            file_path = os.path.join(audio_folder_path, filename_wav)

            # Save the audio recording and limit the files in the folder
            save_recording(file_path)
            limit_audio_files(audio_folder_path, max_audio_files)
            
            audio_file = open(file_path, "rb")
            outtext=""
            l = split_wav_file(file_path)
            for file_path in l:
                audio_file = open(file_path, "rb")
                outtext += self.speech_to_text(audio_file)
                self.text_area1.setText(outtext)
                audio_file.close()
            self.text_area2.setText("SOAP note generation in progress...")
            time.sleep(2)
            soap_notes = self.ai_scribe(0, outtext)
            self.text_area2.setText(soap_notes)
            # Reset the timer label to 0s when recording stops
            self.timer_label.setText("Recording Time: 0s")
            
            #self.toggle_edit_confirm()
            self.populate_picklist()
            
    def toggle_picklist_visibility(self, state):
        if state == Qt.Checked:
            self.populate_picklist()
            self.picklist_container.setVisible(True)
            self.text_area1.setText("Select the audio file from list")
        else:
            if self.text_area2.toPlainText() != "Your SOAP note will be generated here :)":
                self.patient_name_label.setText('')
            self.picklist_container.setVisible(False)
            self.text_area2.setText("Your SOAP note will be generated here :)")
            self.text_area1.setText("Start the recording or select the existing audio!")
        return

            
    def OnClick_Generate(self):
        audio_file = self.wav_file_picklist.currentText()
        if audio_file == "Select audio file":
            QMessageBox.warning(self, "Warning", "Please select the audio file")
            return
        
        file_path = os.path.join(audio_folder_path, audio_file)
        patient_name = audio_file.replace("_", " ").split('-')[0]
        self.patient_name_label.setText(patient_name)
        
        outtext=""
        l = split_wav_file(file_path)
        for file_path in l:
            audio_file = open(file_path, "rb")
            outtext += self.speech_to_text(audio_file)
            self.text_area1.setText(outtext)
            audio_file.close()
        
        self.text_area2.setText(self.ai_scribe(0,outtext))
        
            
    def speech_to_text(self, audio_file, language="en"):
        for i in range(3):
            try:
                transcription = client.audio.transcriptions.create(
                  model="whisper-1", 
                  file=audio_file, 
                  response_format="text",
                  language = language
                )
                return transcription
            except:
                continue
        return 'S.O.A.P Notes cannot be generated at this moment please check your internet connectivity.'

    # Update recording timer
    def update_timer(self):
        global orginal_text, stream, audio_data
        elapsed_time = int(get_audio_length()) 
        min_ = elapsed_time//60
        sec_ = elapsed_time%60
        self.timer_label.setText(f"Recording Time: {min_}:{sec_}")
    
    #gpt-3.5-turbo
    def ai_scribe(self, temperature, content):
        for i in range(3):
            try:
                response = client.chat.completions.create(
                    model="gpt-4o",
                    temperature=temperature,
                    messages=[
                        {
                            "role": "system",
                            "content": self.system_prompt
                        },
                        {
                            "role": "user",
                            "content": content
                        }
                    ]
                )
                return response.choices[0].message.content
            except:
                continue
        return 'S.O.A.P Notes cannot be generated at this moment please check your internet connectivity.'


# Initialize the application
if __name__ == "__main__":
    should_open_app = True
#     if "b6:2d:8f:66:82:00" == get_mac_address():
#         should_open_app = True
    if not should_open_app:
        app = QApplication(sys.argv)
        msg = QMessageBox()
        msg.setIcon(QMessageBox.Warning)
        msg.setText("Application cannot be opened!")
        msg.setInformativeText("This App don't have license to run in current system")
        msg.setWindowTitle("Warning")
        msg.exec_()
        sys.exit()

    app = QApplication(sys.argv)
    window = VoiceRecorderApp()
    window.show()
    sys.exit(app.exec_())
